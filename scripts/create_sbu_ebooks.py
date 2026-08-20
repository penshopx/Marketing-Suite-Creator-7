from __future__ import annotations

from copy import deepcopy
from datetime import date
from html import escape
from pathlib import Path
import re
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from weasyprint import HTML


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "attached_assets" / "ebook-SBU_1787243508753.docx"
OUTPUT = ROOT / "deliverables"
OUTPUT.mkdir(exist_ok=True)

MAIN_TITLE = "SBU & SKK\nTanpa Bingung"
MAIN_SUBTITLE = (
    "Panduan Lengkap Pemenuhan Persyaratan Sertifikasi Konstruksi "
    "di Era Regulasi Baru"
)
MAIN_TAGLINE = "PP 28/2025 • Permen PU 6/2025 • Checklist • Roadmap • Template"
BRAND = "[NAMA BRAND ANDA]"

MINI_TITLE = "Audit SBU–SKK\n14 Hari"
MINI_SUBTITLE = (
    "Mini Ebook Teknis untuk Memetakan Gap, Menata Dokumen, "
    "dan Menentukan Prioritas Kepatuhan"
)
MINI_TAGLINE = "BONUS GRATIS • LEAD MAGNET • PRAKTIS DIGUNAKAN"

INK = "#14213D"
PURPLE = "#5937D8"
VIOLET = "#8B5CF6"
MINT = "#DDF7EE"
GOLD = "#F8C95B"
PAPER = "#FBFAF7"
MUTED = "#5B6475"


def cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell: _Cell, color: str = "D7DCEC") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell: _Cell, text: str, bold: bool = False, color: str = "1B2437") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document, mini: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in (
        ("Title", 29, "14213D"),
        ("Heading 1", 19, "14213D"),
        ("Heading 2", 13, "5937D8"),
        ("Heading 3", 11, "14213D"),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(15 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        ("Bonus Mini Ebook • " if mini else "SBU & SKK Tanpa Bingung • ")
        + "Edisi Digital"
    )
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("7C8494")
    footer.add_run("  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_cover_docx(doc: Document, title: str, subtitle: str, tagline: str, free: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(84 if free else 78)
    r = p.add_run("BONUS GRATIS" if free else "EDISI KOMERSIAL")
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("5937D8")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    for index, line in enumerate(title.split("\n")):
        r = p.add_run(line)
        r.bold = True
        r.font.name = "Aptos Display"
        r.font.size = Pt(35 if not free else 31)
        r.font.color.rgb = RGBColor.from_string("14213D")
        if index < len(title.split("\n")) - 1:
            r.add_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("46546D")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = box.cell(0, 0)
    cell_shading(cell, "EEEAFE")
    cell_border(cell, "C7BEF4")
    set_cell_text(cell, tagline, bold=True, color="5937D8")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150 if not free else 140)
    r = p.add_run(BRAND)
    r.font.name = "Aptos"
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("14213D")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Edisi digital • {date.today().year}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("6B7280")
    doc.add_page_break()


def add_disclaimer_docx(doc: Document, paid: bool = True) -> None:
    doc.add_heading("Catatan Penting Sebelum Anda Memulai", level=1)
    for text in (
        "Ebook ini disusun sebagai materi edukasi dan panduan kerja internal. "
        "Ia tidak menggantikan nasihat hukum, keputusan lembaga berwenang, atau "
        "ketentuan teknis dari OSS, Kementerian PU, LSBU, LSP, LPJK, dan BNSP.",
        "Regulasi, ketentuan peralihan, nomenklatur klasifikasi, persyaratan "
        "dokumen, dan mekanisme layanan dapat berubah. Selalu verifikasi naskah "
        "resmi terbaru sebelum mengambil keputusan atau mengajukan permohonan.",
        "Untuk versi distribusi komersial, ganti placeholder [NAMA BRAND ANDA], "
        "[TAUTAN KONSULTASI / WHATSAPP ANDA], dan [EMAIL ANDA] dengan identitas "
        "bisnis Anda sebelum menjual atau membagikannya.",
    ):
        doc.add_paragraph(text)
    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    cell_shading(cell, "FFF4D6")
    cell_border(cell, "EFC969")
    set_cell_text(
        cell,
        "PRINSIP KERJA: gunakan ebook ini untuk menata bukti, memetakan gap, "
        "dan mengajukan pertanyaan yang tepat kepada kanal resmi.",
        bold=True,
        color="7A5600",
    )
    doc.add_page_break()


def add_contents_docx(doc: Document, items: list[str]) -> None:
    doc.add_heading("Daftar Isi", level=1)
    for number, item in enumerate(items, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"{number:02d}  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string("5937D8")
        p.add_run(item)
    doc.add_page_break()


def iter_document_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def table_to_docx(doc: Document, source: Table) -> None:
    rows = len(source.rows)
    cols = len(source.columns)
    target = doc.add_table(rows=rows, cols=cols)
    target.style = "Table Grid"
    target.autofit = True
    for row_idx, src_row in enumerate(source.rows):
        for col_idx, src_cell in enumerate(src_row.cells):
            target_cell = target.cell(row_idx, col_idx)
            set_cell_text(
                target_cell,
                src_cell.text.replace("\n", " ").strip(),
                bold=row_idx == 0,
                color="FFFFFF" if row_idx == 0 else "1B2437",
            )
            cell_shading(target_cell, "5937D8" if row_idx == 0 else ("F5F6FA" if row_idx % 2 else "FFFFFF"))
            cell_border(target_cell)
        if row_idx == 0:
            set_repeat_table_header(target.rows[row_idx])
    doc.add_paragraph()


def classify_source_paragraph(text: str, chapter_title_pending: bool) -> tuple[str, bool]:
    clean = text.strip()
    if re.fullmatch(r"Bab\s+\d+\s*", clean, flags=re.I):
        return "chapter_number", True
    if chapter_title_pending:
        return "chapter_title", False
    if clean.startswith("Aksi Sekarang"):
        return "action_heading", False
    if clean.startswith("[ ]"):
        return "check_item", False
    if re.match(r"^\d+\.\d+\s", clean):
        return "subheading", False
    if clean in {
        "Tentang Ebook Ini",
        "Kerangka Kerja Buku Ini",
        "Mulai Cepat: Quick-Start & Rencana Aksi 30 Hari",
        "Quick-Start Checklist",
        "Rencana Aksi 30 Hari",
        "Lampiran",
        "Glosarium",
        "Akses langsung (pasang sebagai QR code di versi cetak)",
        "Rujukan regulasi utama",
    }:
        return "heading", False
    if clean.startswith("Bab 8") or clean.startswith("Bab 9"):
        return "chapter_inline", False
    return "body", False


def source_to_docx(doc: Document) -> None:
    source = Document(SOURCE)
    skip_cover = {
        "SBU & SKK Tanpa Bingung :",
        "Panduan Lengkap Pemenuhan Persyaratan Sertifikasi Konstruksi di Era Regulasi Baru",
        "PP 28/2025 dan Permen PU 06 tahun 2025",
    }
    chapter_title_pending = False
    for block in iter_document_blocks(source):
        if isinstance(block, Table):
            table_to_docx(doc, block)
            continue
        text = block.text.strip()
        if not text or text in skip_cover or text.startswith("Judul alternatif:"):
            continue
        kind, chapter_title_pending = classify_source_paragraph(text, chapter_title_pending)
        if kind in {"chapter_title", "chapter_inline"}:
            doc.add_page_break()
            doc.add_heading(text, level=1)
        elif kind == "chapter_number":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(text.upper())
            r.bold = True
            r.font.name = "Aptos"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("5937D8")
        elif kind == "heading":
            doc.add_heading(text, level=1)
        elif kind == "subheading":
            doc.add_heading(text, level=2)
        elif kind == "action_heading":
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell_shading(cell, "DDF7EE")
            cell_border(cell, "92D7BF")
            set_cell_text(cell, text, bold=True, color="0E6B54")
        elif kind == "check_item":
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.12)
            r = p.add_run("□ ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("5937D8")
            p.add_run(text[3:].strip())
        else:
            doc.add_paragraph(text)


def mini_sections() -> list[dict]:
    return [
        {
            "title": "Cara Memakai Mini Ebook Ini",
            "paragraphs": [
                "Panduan ini adalah lembar kerja teknis untuk pemilik, PJT/PJB, "
                "staf legal/perizinan, atau manajer operasional yang perlu "
                "memotret kesiapan SBU–SKK dengan cepat.",
                "Jangan mencoba mengurus semuanya sekaligus. Tarik data, beri "
                "skor, pilih tiga gap paling berisiko, lalu jalankan sprint 14 hari.",
            ],
            "callout": "Target hasil: satu dashboard kepatuhan, tiga prioritas aksi, dan satu pemilik tugas untuk setiap gap.",
        },
        {
            "title": "1. Tarik Data Dasar dalam 30 Menit",
            "paragraphs": [
                "Siapkan satu folder kerja digital bernama “Audit SBU–SKK [BULAN–TAHUN]”. "
                "Jangan mulai dari asumsi atau chat lama. Tarik dokumen versi terakhir yang benar-benar dipakai perusahaan.",
            ],
            "checklist": [
                "SBU aktif, seluruh lampiran, dan tanggal masa berlaku.",
                "NIB, akta pendirian/perubahan, serta daftar KBLI yang terbaca di OSS.",
                "Daftar tenaga ahli/terampil: jabatan kerja, nomor SKK, masa berlaku, dan peran PJT/PJB.",
                "Laporan keuangan terakhir serta status audit KAP bila dipersyaratkan.",
                "Bukti kepemilikan/penguasaan peralatan dan dokumen perjanjian sewa.",
                "Dokumen SMKK/SMK3, kebijakan anti-penyuapan, dan bukti implementasi yang relevan.",
            ],
        },
        {
            "title": "2. Skor Risiko Cepat: Merah, Kuning, atau Hijau",
            "paragraphs": [
                "Nilai setiap item dengan disiplin. Skor ini bukan keputusan regulator; ia adalah alat untuk menentukan pekerjaan mana yang harus dikerjakan lebih dulu.",
            ],
            "table": {
                "headers": ["Skor", "Arti", "Keputusan Operasional"],
                "rows": [
                    ["0 — Merah", "Tidak ada / kedaluwarsa / tidak sinkron", "Tahan pengajuan; tunjuk pemilik aksi hari ini."],
                    ["1 — Kuning", "Ada, tetapi belum lengkap atau belum diverifikasi", "Lengkapi bukti dan tetapkan tenggat maksimal 7 hari."],
                    ["2 — Hijau", "Lengkap, aktif, dan sudah dicek", "Simpan bukti serta jadwalkan pengingat masa berlaku."],
                ],
            },
            "callout": "Aturan prioritas: setiap skor Merah pada KBLI, SBU, SKK PJT/PJB, atau dokumen inti tender harus masuk tiga besar sprint Anda.",
        },
        {
            "title": "3. Matriks Gap Teknis SBU–SKK",
            "paragraphs": [
                "Salin tabel ini ke spreadsheet. Isi berdasarkan dokumen, bukan ingatan. Kolom “bukti” wajib diisi dengan nama file atau tautan folder agar audit dapat ditelusuri.",
            ],
            "table": {
                "headers": ["Area Audit", "Skor 0–2", "Bukti yang Dicek", "Gap / Risiko", "Aksi 14 Hari", "Pemilik"],
                "rows": [
                    ["Akta, NIB, KBLI", "", "", "", "", ""],
                    ["SBU: klasifikasi, subklasifikasi, kualifikasi", "", "", "", "", ""],
                    ["SKK tenaga ahli & tenaga terampil", "", "", "", "", ""],
                    ["PJT / PJB", "", "", "", "", ""],
                    ["Keuangan & audit KAP", "", "", "", "", ""],
                    ["Peralatan utama", "", "", "", "", ""],
                    ["SMKK / SMK3", "", "", "", "", ""],
                    ["SMAP / integritas", "", "", "", "", ""],
                ],
            },
        },
        {
            "title": "4. Audit SKK: Jangan Hanya Cek Tanggal",
            "paragraphs": [
                "SKK yang tampak “masih aktif” belum tentu menjawab kebutuhan SBU atau tender. Cocokkan tiga hal secara bersamaan: jabatan kerja, keterkaitan terhadap kebutuhan badan usaha, dan masa berlaku.",
            ],
            "checklist": [
                "Cocokkan jabatan kerja dan kode dengan kebutuhan klasifikasi/subklasifikasi yang dituju.",
                "Tandai SKK yang habis dalam 12 bulan; jangan menunggu masa berlaku berakhir.",
                "Pastikan PJT dan PJB tidak menjadi titik kegagalan tunggal; siapkan kandidat cadangan bila memungkinkan.",
                "Kumpulkan portofolio, pengalaman proyek, dan bukti pendukung sebelum pendaftaran asesmen.",
                "Verifikasi LSP, TUK, skema, dan informasi uji melalui kanal resmi.",
            ],
        },
        {
            "title": "5. Sprint Kepatuhan 14 Hari",
            "table": {
                "headers": ["Hari", "Fokus", "Output Minimum"],
                "rows": [
                    ["1–2", "Kumpulkan dokumen & buat folder kontrol", "Inventaris dokumen beserta versi dan tanggal."],
                    ["3–4", "Audit legalitas & KBLI", "Daftar ketidaksinkronan akta–NIB–KBLI."],
                    ["5–6", "Audit SBU & subklasifikasi", "Peta status SBU dan pertanyaan untuk LSBU."],
                    ["7–8", "Audit SKK, PJT, PJB", "Kalender SKK dan kebutuhan asesmen/pengganti."],
                    ["9–10", "Audit keuangan & peralatan", "Daftar bukti yang kurang dan penanggung jawabnya."],
                    ["11–12", "SMKK/SMK3 dan integritas", "Bukti kebijakan, pelaksanaan, serta kekurangannya."],
                    ["13", "Rapat keputusan 45 menit", "Tiga gap prioritas, tenggat, dan pemilik tugas."],
                    ["14", "Susun pengingat berulang", "Kalender 30/90/180 hari dan jalur eskalasi."],
                ],
            },
            "callout": "Jangan menyebut sprint selesai sebelum semua pemilik aksi menerima tenggat tertulis dan lokasi folder buktinya.",
        },
        {
            "title": "6. Katalog Bukti: Cara Membuat Audit Tidak Berulang",
            "paragraphs": [
                "Gunakan pola nama file yang konsisten agar orang baru di tim dapat menemukan bukti tanpa mencari di WhatsApp:",
                "YYYY-MM-DD_JenisDokumen_Subjek_Status.pdf",
                "Contoh: 2026-01-15_SKK_PJT-NamaAktif_Aktif.pdf atau 2026-01-20_PerjanjianSewa_Excavator-A01_Terverifikasi.pdf.",
                "Buat satu spreadsheet kontrol dengan kolom: dokumen, pemilik, tanggal terbit, tanggal berakhir, lokasi file, status skor, dan tanggal cek berikutnya.",
            ],
        },
        {
            "title": "7. CTA: Dapatkan Review Gap Anda",
            "paragraphs": [
                "Sudah mengisi matriks? Jangan berhenti di checklist. Kirim ringkasan tiga gap Merah dan tenggat tender/perpanjangan Anda untuk mendapatkan arahan langkah awal.",
                "Konsultasi awal / cek hasil audit: [TAUTAN KONSULTASI / WHATSAPP ANDA]",
                "Email: [EMAIL ANDA]",
                "Nama brand: [NAMA BRAND ANDA]",
            ],
            "callout": "Lead magnet ini dirancang untuk membuka percakapan yang lebih berkualitas: calon klien datang dengan data, bukan hanya pertanyaan umum.",
        },
    ]


def add_mini_docx_content(doc: Document) -> None:
    for index, section in enumerate(mini_sections()):
        if index > 0:
            doc.add_page_break()
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paragraphs", []):
            p = doc.add_paragraph(paragraph)
            if paragraph.startswith("YYYY-MM-DD_"):
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
        for item in section.get("checklist", []):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.12)
            r = p.add_run("□ ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("5937D8")
            p.add_run(item)
        table_data = section.get("table")
        if table_data:
            headers = table_data["headers"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for i, value in enumerate(headers):
                set_cell_text(table.cell(0, i), value, bold=True, color="FFFFFF")
                cell_shading(table.cell(0, i), "5937D8")
                cell_border(table.cell(0, i))
            set_repeat_table_header(table.rows[0])
            for row_idx, row in enumerate(table_data["rows"]):
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    set_cell_text(cells[i], value)
                    cell_shading(cells[i], "F7F6FC" if row_idx % 2 == 0 else "FFFFFF")
                    cell_border(cells[i])
        if section.get("callout"):
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell_shading(cell, "DDF7EE")
            cell_border(cell, "92D7BF")
            set_cell_text(cell, section["callout"], bold=True, color="0E6B54")


def render_table_html(table: Table) -> str:
    rows = []
    for row_idx, row in enumerate(table.rows):
        tag = "th" if row_idx == 0 else "td"
        cells = "".join(f"<{tag}>{escape(cell.text.replace(chr(10), ' ').strip())}</{tag}>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return f"<div class='table-wrap'><table>{''.join(rows)}</table></div>"


def source_to_html() -> str:
    source = Document(SOURCE)
    skip_cover = {
        "SBU & SKK Tanpa Bingung :",
        "Panduan Lengkap Pemenuhan Persyaratan Sertifikasi Konstruksi di Era Regulasi Baru",
        "PP 28/2025 dan Permen PU 06 tahun 2025",
    }
    parts = []
    chapter_title_pending = False
    for block in iter_document_blocks(source):
        if isinstance(block, Table):
            parts.append(render_table_html(block))
            continue
        text = block.text.strip()
        if not text or text in skip_cover or text.startswith("Judul alternatif:"):
            continue
        kind, chapter_title_pending = classify_source_paragraph(text, chapter_title_pending)
        safe = escape(text)
        if kind in {"chapter_title", "chapter_inline"}:
            parts.append(f"<h1 class='chapter'>{safe}</h1>")
        elif kind == "chapter_number":
            parts.append(f"<p class='chapter-no'>{safe.upper()}</p>")
        elif kind == "heading":
            parts.append(f"<h1>{safe}</h1>")
        elif kind == "subheading":
            parts.append(f"<h2>{safe}</h2>")
        elif kind == "action_heading":
            parts.append(f"<div class='action'><strong>{safe}</strong></div>")
        elif kind == "check_item":
            parts.append(f"<p class='check'>□ {escape(text[3:].strip())}</p>")
        else:
            parts.append(f"<p>{safe}</p>")
    return "\n".join(parts)


def render_mini_html() -> str:
    parts = []
    for section in mini_sections():
        parts.append(f"<section class='mini-section'><h1>{escape(section['title'])}</h1>")
        for paragraph in section.get("paragraphs", []):
            klass = "code-line" if paragraph.startswith("YYYY-MM-DD_") else ""
            parts.append(f"<p class='{klass}'>{escape(paragraph)}</p>")
        for item in section.get("checklist", []):
            parts.append(f"<p class='check'>□ {escape(item)}</p>")
        if section.get("table"):
            data = section["table"]
            parts.append("<div class='table-wrap'><table><tr>" + "".join(
                f"<th>{escape(value)}</th>" for value in data["headers"]
            ) + "</tr>")
            for row in data["rows"]:
                parts.append("<tr>" + "".join(f"<td>{escape(value)}</td>" for value in row) + "</tr>")
            parts.append("</table></div>")
        if section.get("callout"):
            parts.append(f"<div class='action'><strong>{escape(section['callout'])}</strong></div>")
        parts.append("</section>")
    return "\n".join(parts)


def pdf_html(title: str, subtitle: str, tagline: str, body: str, free: bool) -> str:
    label = "BONUS GRATIS" if free else "EDISI KOMERSIAL"
    footer = "Bonus Mini Ebook • Audit SBU–SKK 14 Hari" if free else "SBU & SKK Tanpa Bingung"
    return f"""<!doctype html>
<html lang="id">
<head>
<meta charset="utf-8">
<style>
@page {{
  size: A4;
  margin: 19mm 17mm 17mm;
  @bottom-center {{
    content: "{footer}  •  " counter(page);
    color: #737B8C; font: 8.5pt "DejaVu Sans", sans-serif;
  }}
}}
* {{ box-sizing: border-box; }}
body {{ color: {INK}; background: white; font-family: "DejaVu Sans", Arial, sans-serif; font-size: 10.1pt; line-height: 1.52; }}
.cover {{ min-height: 258mm; display: flex; flex-direction: column; justify-content: center; align-items: center; text-align: center; page-break-after: always; background: linear-gradient(145deg, #F9F8FF 0%, #FFFFFF 48%, #E8FFF5 100%); border: 2px solid #E6E2FC; padding: 25mm 14mm; }}
.eyebrow {{ color: {PURPLE}; font-weight: 700; letter-spacing: 1.7px; font-size: 9pt; }}
.cover h1 {{ margin: 10mm 0 5mm; font-family: "DejaVu Sans", Arial, sans-serif; font-size: {"35pt" if not free else "30pt"}; line-height: 1.04; color: {INK}; }}
.cover .subtitle {{ max-width: 138mm; margin: 0; color: #46546D; font-size: 14pt; line-height: 1.35; }}
.pill {{ margin: 13mm 0 0; padding: 4.5mm 6mm; background: #EEEAFE; border: 1px solid #C9BFF2; border-radius: 4mm; color: {PURPLE}; font-weight: 700; font-size: 9.5pt; }}
.brand {{ margin-top: auto; color: {INK}; font-weight: 700; font-size: 11pt; }}
.edition {{ color: #687386; font-size: 8.5pt; margin: 2mm 0 0; }}
.legal {{ page-break-after: always; }}
h1 {{ color: {INK}; font-size: 19pt; line-height: 1.18; margin: 15mm 0 5mm; page-break-after: avoid; }}
h1.chapter {{ page-break-before: always; }}
h2 {{ color: {PURPLE}; font-size: 12.5pt; line-height: 1.25; margin: 8mm 0 3mm; page-break-after: avoid; }}
p {{ margin: 0 0 3.5mm; }}
.chapter-no {{ color: {PURPLE}; font-weight: 700; letter-spacing: 1px; font-size: 9pt; margin: 8mm 0 -2mm; }}
.action {{ background: {MINT}; border-left: 4px solid #26A37F; padding: 4mm 5mm; margin: 5mm 0; color: #0A614A; page-break-inside: avoid; }}
.check {{ margin-left: 3mm; }}
.table-wrap {{ overflow: hidden; margin: 5mm 0 6mm; page-break-inside: avoid; }}
table {{ width: 100%; border-collapse: collapse; font-size: 8.25pt; line-height: 1.3; }}
th {{ background: {PURPLE}; color: white; font-weight: 700; text-align: left; }}
th, td {{ border: 1px solid #D8DCEA; padding: 2.2mm; vertical-align: top; }}
tr:nth-child(even) td {{ background: #F7F8FC; }}
.toc {{ page-break-after: always; }}
.toc ol {{ padding-left: 0; list-style: none; }}
.toc li {{ border-bottom: 1px solid #E5E7EF; padding: 2.7mm 0; }}
.toc-no {{ color: {PURPLE}; font-weight: 700; display: inline-block; width: 12mm; }}
.code-line {{ font-family: "DejaVu Sans Mono", monospace; font-size: 8.8pt; background: #F4F5F9; padding: 2.5mm; }}
.mini-section {{ page-break-before: always; }}
.mini-section:first-of-type {{ page-break-before: auto; }}
</style>
</head>
<body>
<section class="cover">
  <div class="eyebrow">{label}</div>
  <h1>{title.replace(chr(10), '<br>')}</h1>
  <p class="subtitle">{subtitle}</p>
  <div class="pill">{tagline}</div>
  <div class="brand">{BRAND}</div>
  <p class="edition">Edisi digital • {date.today().year}</p>
</section>
<section class="legal">
  <h1>Catatan Penting Sebelum Anda Memulai</h1>
  <p>Ebook ini disusun sebagai materi edukasi dan panduan kerja internal. Ia tidak menggantikan nasihat hukum, keputusan lembaga berwenang, atau ketentuan teknis dari OSS, Kementerian PU, LSBU, LSP, LPJK, dan BNSP.</p>
  <p>Regulasi, ketentuan peralihan, nomenklatur klasifikasi, persyaratan dokumen, dan mekanisme layanan dapat berubah. Selalu verifikasi naskah resmi terbaru sebelum mengambil keputusan atau mengajukan permohonan.</p>
  <p>Untuk versi distribusi, ganti placeholder <strong>[NAMA BRAND ANDA]</strong>, <strong>[TAUTAN KONSULTASI / WHATSAPP ANDA]</strong>, dan <strong>[EMAIL ANDA]</strong> dengan identitas bisnis Anda.</p>
  <div class="action"><strong>PRINSIP KERJA: gunakan ebook ini untuk menata bukti, memetakan gap, dan mengajukan pertanyaan yang tepat kepada kanal resmi.</strong></div>
</section>
{body}
</body></html>"""


def create_main_docx() -> Path:
    doc = Document()
    style_document(doc)
    add_cover_docx(doc, MAIN_TITLE, MAIN_SUBTITLE, MAIN_TAGLINE, free=False)
    add_disclaimer_docx(doc)
    add_contents_docx(
        doc,
        [
            "Tentang Ebook Ini",
            "Mulai Cepat: Quick-Start & Rencana Aksi 30 Hari",
            "Bab 1 — Ketika Aturan Berubah di Tengah Jalan",
            "Bab 2 — Peta Regulasi: Membaca Ulang Medan",
            "Bab 3 — SBU: Persyaratan Baru Badan Usaha",
            "Bab 4 — SKK: Tiket Karier Tenaga Kerja Konstruksi",
            "Bab 5 — Gap Analysis & Transisi",
            "Bab 6 — Checklist & Template Siap Pakai",
            "Bab 7 — Studi Kasus",
            "Bab 8 — FAQ",
            "Bab 9 — Kepatuhan sebagai Senjata Kompetitif",
            "Lampiran & Rujukan",
        ],
    )
    source_to_docx(doc)
    path = OUTPUT / "Ebook_SBU_SKK_Tanpa_Bingung_Edisi_Komersial.docx"
    doc.save(path)
    return path


def create_mini_docx() -> Path:
    doc = Document()
    style_document(doc, mini=True)
    add_cover_docx(doc, MINI_TITLE, MINI_SUBTITLE, MINI_TAGLINE, free=True)
    add_disclaimer_docx(doc, paid=False)
    add_contents_docx(doc, [section["title"] for section in mini_sections()])
    add_mini_docx_content(doc)
    path = OUTPUT / "Bonus_Mini_Ebook_Audit_SBU_SKK_14_Hari.docx"
    doc.save(path)
    return path


def main() -> None:
    main_docx = create_main_docx()
    mini_docx = create_mini_docx()
    main_pdf = OUTPUT / "Ebook_SBU_SKK_Tanpa_Bingung_Edisi_Komersial.pdf"
    mini_pdf = OUTPUT / "Bonus_Mini_Ebook_Audit_SBU_SKK_14_Hari.pdf"
    main_body = (
        "<section class='toc'><h1>Daftar Isi</h1><ol>"
        + "".join(
            f"<li><span class='toc-no'>{i:02d}</span>{escape(item)}</li>"
            for i, item in enumerate(
                [
                    "Tentang Ebook Ini",
                    "Mulai Cepat: Quick-Start & Rencana Aksi 30 Hari",
                    "Bab 1 — Ketika Aturan Berubah di Tengah Jalan",
                    "Bab 2 — Peta Regulasi: Membaca Ulang Medan",
                    "Bab 3 — SBU: Persyaratan Baru Badan Usaha",
                    "Bab 4 — SKK: Tiket Karier Tenaga Kerja Konstruksi",
                    "Bab 5 — Gap Analysis & Transisi",
                    "Bab 6 — Checklist & Template Siap Pakai",
                    "Bab 7 — Studi Kasus",
                    "Bab 8 — FAQ",
                    "Bab 9 — Kepatuhan sebagai Senjata Kompetitif",
                    "Lampiran & Rujukan",
                ],
                1,
            )
        )
        + "</ol></section>"
        + source_to_html()
    )
    HTML(string=pdf_html(MAIN_TITLE, MAIN_SUBTITLE, MAIN_TAGLINE, main_body, free=False)).write_pdf(main_pdf)
    HTML(string=pdf_html(MINI_TITLE, MINI_SUBTITLE, MINI_TAGLINE, render_mini_html(), free=True)).write_pdf(mini_pdf)

    bundle = OUTPUT / "Paket_Ebook_SBU_SKK_Siap_Jual.zip"
    with zipfile.ZipFile(bundle, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in (main_pdf, main_docx, mini_pdf, mini_docx):
            archive.write(path, arcname=path.name)
    print(f"Created: {main_pdf.name}, {main_docx.name}, {mini_pdf.name}, {mini_docx.name}, {bundle.name}")


if __name__ == "__main__":
    main()